import os
import io
from typing import Optional
from sqlalchemy.orm import Session
from sqlalchemy import func, desc, extract
from datetime import datetime, timedelta
from collections import Counter

from app.models.models import Story, Chapter, User, Language, Genre, Emotion, StoryLength, Gender
from app.schemas.schemas import StoryCreate, StoriesPageResponse, AnalyticsResponse
from app.services.ai_service import story_generator


def count_words(text: str) -> int:
    return len(text.split())


def calc_read_time(word_count: int) -> float:
    return round(word_count / 200, 1)  # avg 200 wpm


async def create_story(db: Session, user: User, payload: StoryCreate) -> Story:
    """Generate and persist a new story."""
    story_data = await story_generator.generate_story(
        genre=payload.genre,
        emotion=payload.emotion,
        story_length=payload.story_length,
        language=payload.language,
        protagonist_name=payload.protagonist_name or user.full_name,
        protagonist_age=payload.protagonist_age or user.age,
        protagonist_gender=payload.protagonist_gender or user.gender,
        protagonist_interests=payload.protagonist_interests or user.interests,
        custom_topic=payload.custom_topic,
    )

    # Build full content string
    full_content = "\n\n".join(
        "## {}\n\n{}".format(
            ch.get('title', 'Chapter {}'.format(ch['chapter_number'])),
            ch['content']
        )
        for ch in story_data.get("chapters", [])
    )

    word_count = count_words(full_content)
    read_time = calc_read_time(word_count)

    story = Story(
        user_id=user.id,
        title=story_data.get("title", "Untitled Story"),
        content=full_content,
        moral=story_data.get("moral"),
        language=payload.language,
        genre=payload.genre,
        emotion=payload.emotion,
        story_length=payload.story_length,
        protagonist_name=payload.protagonist_name,
        protagonist_age=payload.protagonist_age,
        protagonist_gender=payload.protagonist_gender,
        protagonist_interests=payload.protagonist_interests,
        word_count=word_count,
        read_time_minutes=read_time,
        model_used=story_data.get("model_used"),
        generation_time_seconds=story_data.get("generation_time_seconds"),
    )
    db.add(story)
    db.flush()

    for ch_data in story_data.get("chapters", []):
        chapter = Chapter(
            story_id=story.id,
            chapter_number=ch_data.get("chapter_number", 1),
            title=ch_data.get("title"),
            content=ch_data.get("content", ""),
            image_prompt=ch_data.get("image_prompt"),
        )
        db.add(chapter)

    db.commit()
    db.refresh(story)
    return story


def get_user_stories(
    db: Session,
    user_id: str,
    page: int = 1,
    page_size: int = 12,
    search: Optional[str] = None,
    genre: Optional[Genre] = None,
    language: Optional[Language] = None,
    emotion: Optional[Emotion] = None,
    favorites_only: bool = False,
) -> StoriesPageResponse:
    query = db.query(Story).filter(Story.user_id == user_id)

    if search:
        query = query.filter(Story.title.ilike(f"%{search}%"))
    if genre:
        query = query.filter(Story.genre == genre)
    if language:
        query = query.filter(Story.language == language)
    if emotion:
        query = query.filter(Story.emotion == emotion)
    if favorites_only:
        query = query.filter(Story.is_favorite == True)

    total = query.count()
    stories = (
        query.order_by(desc(Story.created_at))
        .offset((page - 1) * page_size)
        .limit(page_size)
        .all()
    )

    return StoriesPageResponse(
        stories=stories,
        total=total,
        page=page,
        page_size=page_size,
        total_pages=(total + page_size - 1) // page_size,
    )


def get_story_by_id(db: Session, story_id: str, user_id: str) -> Optional[Story]:
    return db.query(Story).filter(
        Story.id == story_id, Story.user_id == user_id
    ).first()


def toggle_favorite(db: Session, story_id: str, user_id: str) -> Story:
    story = get_story_by_id(db, story_id, user_id)
    if not story:
        raise ValueError("Story not found")
    story.is_favorite = not story.is_favorite
    db.commit()
    db.refresh(story)
    return story


def delete_story(db: Session, story_id: str, user_id: str) -> bool:
    story = get_story_by_id(db, story_id, user_id)
    if not story:
        return False
    db.delete(story)
    db.commit()
    return True


def get_analytics(db: Session, user_id: str) -> AnalyticsResponse:
    stories = db.query(Story).filter(Story.user_id == user_id).all()

    total_stories = len(stories)
    total_words = sum(s.word_count or 0 for s in stories)
    favorites = sum(1 for s in stories if s.is_favorite)

    genre_counts = Counter(s.genre.value for s in stories)
    language_counts = Counter(s.language.value for s in stories)
    emotion_counts = Counter(s.emotion.value for s in stories)

    fav_genre = genre_counts.most_common(1)[0][0] if genre_counts else None
    fav_language = language_counts.most_common(1)[0][0] if language_counts else None

    # Monthly distribution (last 12 months)
    monthly = {}
    for story in stories:
        if story.created_at:
            key = story.created_at.strftime("%Y-%m")
            monthly[key] = monthly.get(key, 0) + 1

    monthly_list = [{"month": k, "count": v} for k, v in sorted(monthly.items())[-12:]]

    avg_words = total_words / total_stories if total_stories > 0 else 0
    total_read_time = sum(s.read_time_minutes or 0 for s in stories) / 60

    return AnalyticsResponse(
        total_stories=total_stories,
        total_words=total_words,
        favorite_stories=favorites,
        favorite_genre=fav_genre,
        favorite_language=fav_language,
        genre_distribution=dict(genre_counts),
        language_distribution=dict(language_counts),
        emotion_distribution=dict(emotion_counts),
        monthly_stories=monthly_list,
        avg_word_count=round(avg_words),
        total_read_time_hours=round(total_read_time, 1),
    )


def export_story_txt(story: Story) -> bytes:
    lines = [
        f"STORYVERSE AI — {story.title}",
        "=" * 60,
        f"Genre: {story.genre.value.title()} | Language: {story.language.value.title()}",
        f"Emotion: {story.emotion.value.title()} | Length: {story.story_length.value.title()}",
        f"Generated: {story.created_at.strftime('%B %d, %Y')}",
        "=" * 60,
        "",
        story.content,
        "",
        "=" * 60,
        f"MORAL: {story.moral or 'N/A'}",
        "=" * 60,
        "",
        "Generated by StoryVerse AI — storyverse.ai",
    ]
    return "\n".join(lines).encode("utf-8")


def export_story_pdf(story: Story) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "StoryTitle", parent=styles["Title"],
        fontSize=24, spaceAfter=6, textColor=colors.HexColor("#6C3FC5"),
        alignment=TA_CENTER,
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#888888"),
        alignment=TA_CENTER, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=11, leading=18, spaceAfter=12,
        alignment=TA_JUSTIFY,
    )
    chapter_style = ParagraphStyle(
        "Chapter", parent=styles["Heading2"],
        fontSize=14, textColor=colors.HexColor("#6C3FC5"),
        spaceBefore=16, spaceAfter=8,
    )
    moral_style = ParagraphStyle(
        "Moral", parent=styles["Normal"],
        fontSize=11, textColor=colors.HexColor("#444"), italic=True,
        borderPad=8, backColor=colors.HexColor("#F5F0FF"),
    )

    content = []
    content.append(Paragraph(story.title, title_style))
    content.append(Paragraph(
        f"{story.genre.value.title()} · {story.language.value.title()} · {story.emotion.value.title()}",
        meta_style
    ))
    content.append(Paragraph(
        f"Generated on {story.created_at.strftime('%B %d, %Y')} · {story.word_count} words",
        meta_style
    ))
    content.append(HRFlowable(color=colors.HexColor("#6C3FC5"), thickness=1, spaceAfter=12))

    # Chapters
    for chapter in story.chapters:
        if chapter.title:
            content.append(Paragraph(f"Chapter {chapter.chapter_number}: {chapter.title}", chapter_style))
        for para in chapter.content.split("\n\n"):
            if para.strip():
                content.append(Paragraph(para.strip(), body_style))

    if story.moral:
        content.append(Spacer(1, 12))
        content.append(HRFlowable(color=colors.HexColor("#6C3FC5"), thickness=1, spaceAfter=8))
        content.append(Paragraph(f"<b>Moral:</b> {story.moral}", moral_style))

    content.append(Spacer(1, 20))
    content.append(Paragraph("Generated by StoryVerse AI", meta_style))

    doc.build(content)
    return buffer.getvalue()


def export_story_docx(story: Story) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(story.title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x6C, 0x3F, 0xC5)

    # Metadata
    meta = doc.add_paragraph(
        f"{story.genre.value.title()} · {story.language.value.title()} · "
        f"{story.emotion.value.title()} · {story.created_at.strftime('%B %d, %Y')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    for chapter in story.chapters:
        doc.add_heading(
            f"Chapter {chapter.chapter_number}: {chapter.title or ''}",
            level=1
        )
        for para_text in chapter.content.split("\n\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(8)

    if story.moral:
        doc.add_paragraph()
        moral_para = doc.add_paragraph()
        run = moral_para.add_run(f"Moral: {story.moral}")
        run.bold = True
        run.font.color.rgb = RGBColor(0x6C, 0x3F, 0xC5)

    doc.add_paragraph()
    footer = doc.add_paragraph("Generated by StoryVerse AI")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
